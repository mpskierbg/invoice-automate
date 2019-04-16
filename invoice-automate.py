import os
import docx
from docx.shared import Pt
from docx import Document
from docx.shared import Length
from datetime import date, timedelta
from docx.shared import Inches
from docx.shared import RGBColor
import win32com.client

def sendmail(attachment):
    today = date.today().strftime("%B %d, %Y")
    o = win32com.client.Dispatch("Outlook.Application")

    Msg = o.CreateItem(0)
    Msg.Importance = 0
    Msg.Subject = today +' Invoice'
    Msg.HTMLBody = '<p> Hello Leslie, <br> <br> Here is the invoice for ' + today + '<br><br>In Solidarity,<br>Matthew Senko<p/>'

    Msg.To = "matthew.senko@allensuperiorcourt.us"
    Msg.Attachments.Add("C:\\Users\\MASEAB\\Documents\\" + attachment)


    Msg.Send()
   

def space(doc):
    space = doc.add_paragraph()
    space.paragraph_format.space_after = Pt(0)
    space.paragraph_format.line_spacing = 1.15
    return space

def line(doc):
    line = doc.add_paragraph().add_run("--------------------------------------------------------------------------------------------------------------------------------------------------------------------")
    font = line.font
    font.color.rgb = RGBColor(211,211,211)
    return line

def main(): 

    today = date.today().strftime("%B %d, %Y")
    two = date.today() - timedelta(days=14)
    twoweeks = two.strftime("%B %d, %Y")
    doc = docx.Document()
   
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
   
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.15   
   
    sections = doc.sections

    section = sections[0]
    for section in sections:
        section.left_margin = Inches(.45)
        section.right_margin = Inches(.45)
        section.top_margin = Inches(.45)
        section.bottom_margin = Inches(.45)

    doc.add_paragraph('Matthew Senko')

    doc.add_paragraph('Law Clerk for the Commercial Court of Northern Indiana')

    space(doc)

    doc.add_paragraph('4525 Tacoma Ave')

    doc.add_paragraph('Fort Wayne, IN 46807')

    space(doc)

    doc.add_paragraph('260-450-9556')

    doc.add_paragraph('msenko2@gmail.com')

    space(doc)
    space(doc)
    space(doc)

    doc.add_paragraph(today).alignment = 1

    doc.add_paragraph("Invoice").alignment = 1

    line(doc)

    space(doc)

    doc.add_paragraph('Commercial Courts of Northern Indiana - Allen, Elkhart, and Lake Counties')

    space(doc)

    re = doc.add_paragraph().add_run('RE: Law Clerk services')
    font = re.font
    font.bold = True
    

    line(doc)

    space(doc)

    doc.add_paragraph("Current Charges")
     
    space(doc)

    doc.add_paragraph("Law Clerk Services for the time period of:")

    doc.add_paragraph(twoweeks + " to " + today)

    space(doc)

    doc.add_paragraph("Total fees:")

    doc.add_paragraph("$2,353.85")
    
    doc.save(today + 'invoice.docx')
    os.startfile(today + 'invoice.docx')
    sendmail(today+'invoice.docx')
if __name__ == "__main__":
    main()
